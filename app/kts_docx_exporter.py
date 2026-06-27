"""DOCX export for reviewed SPA/SHA KTS results."""

from __future__ import annotations

import re
from io import BytesIO
from typing import Any

from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from source_refs import clean_clause_ref

EAST_ASIA_FONT = "宋体"
TITLE_FONT = "黑体"
GROUP_ORDER = {"SPA": 0, "SHA": 1}
SKIP_EMPTY_OUTPUT_CATEGORIES = {"conditional_output", "optional_conditional_output"}
PARENTHETICAL_MARKER_RE = re.compile(r"\s*([（(][一二三四五六七八九十\d]+[）)])")
BRACKETED_NOTE_RE = re.compile(r"【[^】]{1,1200}】")
NOTE_LINE_RE = re.compile(r"\s*(【[^】]*注[：:][^】]*】)")
PRENUMBERED_LINE_RE = re.compile(r"^(\d+[.、]\s*|（[一二三四五六七八九十\d]+）|\([0-9]+\))")
SUB_NUMBERED_LINE_RE = re.compile(r"^(\d+\.\d+)")
SOURCE_REF_MAX_LENGTH = 96
SOURCE_REF_MAX_COUNT = 5
INDENT_PER_LEVEL = Cm(0.5)


class KtsDocxExportError(ValueError):
    """Raised when KTS results cannot be exported."""


def protect_bracketed_notes(text: str) -> tuple[str, dict[str, str]]:
    notes: dict[str, str] = {}

    def replace(match: re.Match[str]) -> str:
        token = f"@@KTS_NOTE_{len(notes)}@@"
        notes[token] = match.group(0)
        return token

    return BRACKETED_NOTE_RE.sub(replace, text), notes


def restore_bracketed_notes(text: str, notes: dict[str, str]) -> str:
    restored = text
    for token, note in notes.items():
        restored = restored.replace(token, note)
    return restored


def separate_note_lines(text: str) -> str:
    return NOTE_LINE_RE.sub(r"\n\1", text)


def is_note_line(line: str) -> bool:
    text = line.strip()
    return text.startswith("【") and text.endswith("】") and "注" in text


def split_readable_lines(value: object) -> list[str]:
    text = str(value or "").strip()
    if not text:
        return []
    if "\n" in text:
        prepared = text
    else:
        prepared = re.sub(r"\s+", " ", text)
        prepared, bracketed_notes = protect_bracketed_notes(prepared)
        prepared = re.sub(r"([。；;])\s*", r"\1\n", prepared)
        prepared = re.sub(
            r"([：:])\s*((?:（?[一二三四五六七八九十\d]+[）.)、]))",
            r"\1\n\2",
            prepared,
        )
        prepared = split_parenthetical_markers(prepared)
        prepared = re.sub(r"(其中[，,])\s*", r"\n\1", prepared)
        prepared = restore_bracketed_notes(prepared, bracketed_notes)
        prepared = separate_note_lines(prepared)
    return [line.strip() for line in prepared.split("\n") if line.strip()]


def split_parenthetical_markers(text: str) -> str:
    parts: list[str] = []
    cursor = 0
    for match in PARENTHETICAL_MARKER_RE.finditer(text):
        before = text[: match.start()].rstrip()
        previous_char = before[-1:] if before else ""
        should_split = not before or previous_char in "：:；;。！？!?、\n"
        if not should_split:
            continue
        parts.append(text[cursor : match.start()].rstrip())
        parts.append("\n")
        parts.append(match.group(1))
        cursor = match.end()
    parts.append(text[cursor:])
    return "".join(parts)


def number_readable_lines(lines: list[str]) -> list[str]:
    """Number content lines with two-level scheme: 1. for top-level, (1) for sub-items.

    Lines with label:value structure are top-level items (numbered 1. 2. 3.).
    Consecutive non-label lines (2+) following a top-level item are sub-items,
    numbered (1) (2) (3). A single non-label line after a top-level item is
    just a continuation and gets its own top-level number.
    """
    if len(lines) <= 1:
        return lines

    LABEL_VALUE_RE = re.compile(r"^.{2,15}[：:]")

    # First pass: classify each line
    roles: list[str] = []
    for line in lines:
        if is_note_line(line):
            roles.append("note")
        elif LABEL_VALUE_RE.match(line):
            roles.append("label")
        else:
            roles.append("other")

    # Second pass: determine which "other" runs are sub-item lists (2+ consecutive)
    is_sub_item: list[bool] = [False] * len(lines)
    i = 0
    while i < len(lines):
        if roles[i] == "other":
            run_start = i
            while i < len(lines) and roles[i] == "other":
                i += 1
            if i - run_start >= 2:
                for j in range(run_start, i):
                    is_sub_item[j] = True
        else:
            i += 1

    # Third pass: number
    numbered: list[str] = []
    top_number = 1
    sub_number = 0

    for idx, line in enumerate(lines):
        if roles[idx] == "note":
            numbered.append(line)
            continue

        if roles[idx] == "label":
            numbered.append(f"{top_number}. {line}")
            top_number += 1
            sub_number = 0
        elif is_sub_item[idx]:
            sub_number += 1
            stripped = PRENUMBERED_LINE_RE.sub("", line).strip()
            numbered.append(f"({sub_number}) {stripped}")
        else:
            numbered.append(f"{top_number}. {line}")
            top_number += 1
            sub_number = 0

    return numbered

NOT_FOUND_PATTERNS = re.compile(r"(未见|未载明|待确认|未明确约定|候选证据未见).{0,20}(约定|确认|载明|明确)?")


def _is_not_found_line(line: str) -> bool:
    """Check if a line is purely a not-found placeholder with no substantive content."""
    stripped = line.strip()
    if not stripped:
        return True
    # Remove label prefix (e.g. "签署方：") to check the content part
    content_part = re.sub(r"^.{2,15}[：:]", "", stripped).strip()
    if not content_part:
        return True
    # If the content part is dominated by not-found language
    if NOT_FOUND_PATTERNS.search(content_part):
        clean = NOT_FOUND_PATTERNS.sub("", content_part).strip()
        clean = re.sub(r"[，。；、:：\s]+", " ", clean).strip()
        # Keep only if remaining text has substantive facts (numbers, names, dates)
        has_facts = bool(re.search(r"[\d%％万亿元年月日]", clean))
        if not has_facts and len(clean) < 15:
            return True
    return False


def format_kts_content(value: object) -> list[str]:
    lines = split_readable_lines(value)
    lines = [line for line in lines if not _is_not_found_line(line)]
    return number_readable_lines(lines)


def saved_human_review(item: dict[str, Any]) -> dict[str, Any] | None:
    review = item.get("human_review")
    if not isinstance(review, dict):
        return None
    if any(key in review for key in ("status", "content", "updated_at")):
        return review
    return None


def export_content_lines(item: dict[str, Any]) -> list[str]:
    review = saved_human_review(item)
    if review is not None:
        content = str(review.get("content") or "").strip()
        return split_readable_lines(content)

    draft_content = str(item.get("draft_content") or "").strip()
    if draft_content:
        return format_kts_content(draft_content)
    if str(item.get("status") or "") == "unclear":
        return ["未见明确约定。"]
    return []


def export_label(item: dict[str, Any]) -> str:
    return str(item.get("label") or item.get("taxonomy_id") or "未命名事项").strip()


def should_skip_empty_export_item(item: dict[str, Any]) -> bool:
    output_policy = item.get("output_policy", {})
    if not isinstance(output_policy, dict):
        return False
    return str(output_policy.get("category") or "") in SKIP_EMPTY_OUTPUT_CATEGORIES


CLAUSE_NUMBER_RE = re.compile(r"^(第?[\d一二三四五六七八九十百]+[.、．条][\d.]*\s*)")


def extract_clause_number(value: object) -> str:
    """Extract only the clause number prefix from a ref string."""
    text = str(value or "").strip()
    m = CLAUSE_NUMBER_RE.match(text)
    if m:
        return m.group(1).strip()
    # If the whole string is short and looks like a number ref, keep it
    if re.match(r"^[\d.]+$", text) and len(text) <= 12:
        return text
    return ""


def export_source_refs(item: dict[str, Any]) -> list[str]:
    """Extract only clause numbers for source refs — no prose content."""
    refs = item.get("clause_refs", [])
    if not isinstance(refs, list):
        refs = []

    clean_refs: list[str] = []
    seen: set[str] = set()
    for ref in refs:
        number = extract_clause_number(ref)
        if not number or number in seen:
            continue
        seen.add(number)
        clean_refs.append(number)
        if len(clean_refs) >= SOURCE_REF_MAX_COUNT:
            break
    return clean_refs


def export_items(record: dict[str, Any]) -> list[dict[str, Any]]:
    raw_items = record.get("items", [])
    if not isinstance(raw_items, list):
        raise KtsDocxExportError("KTS 结果格式不正确。")

    items: list[dict[str, Any]] = []
    for index, item in enumerate(raw_items):
        if not isinstance(item, dict):
            continue
        content_lines = export_content_lines(item)
        if not content_lines and should_skip_empty_export_item(item):
            continue
        group = str(item.get("group") or "其他").strip() or "其他"
        items.append(
            {
                "index": index,
                "group": group,
                "label": export_label(item),
                "content_lines": content_lines,
                "source_refs": export_source_refs(item),
            }
        )
    if not items:
        raise KtsDocxExportError("暂无可导出的 KTS 事项。")
    return sorted(items, key=lambda item: (GROUP_ORDER.get(str(item["group"]), 99), item["index"]))


def set_cell_text(cell, text: str, bold: bool = False, align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = EAST_ASIA_FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)


def _line_indent_level(line: str) -> int:
    """Determine indentation level based on numbering pattern."""
    stripped = line.strip()
    if re.match(r"^\d+\.\d+\.\d+", stripped):
        return 2
    if re.match(r"^\d+\.\d+", stripped):
        return 1
    if re.match(r"^[（(][一二三四五六七八九十\d]+[）)]", stripped):
        return 1
    if re.match(r"^\([ivxl]+\)", stripped):
        return 2
    return 0


def append_cell_lines(cell, lines: list[str]) -> None:
    cell.text = ""
    if not lines:
        cell.paragraphs[0].text = ""
        return
    for index, line in enumerate(lines):
        paragraph = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.15
        indent_level = _line_indent_level(line)
        if indent_level > 0:
            paragraph.paragraph_format.left_indent = Cm(0.5 * indent_level)
        run = paragraph.add_run(line)
        run.font.name = EAST_ASIA_FONT
        run.font.size = Pt(10)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)


def append_source_refs(cell, refs: list[str]) -> None:
    if not refs:
        return
    paragraph = cell.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.05
    run = paragraph.add_run("信息来源：" + "  ".join(refs))
    run.font.name = EAST_ASIA_FONT
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(104, 112, 105)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_width(cell, width_cm: float) -> None:
    cell.width = Cm(width_cm)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "B7B7B7")


def set_table_dimensions(table, widths: list[float]) -> None:
    total_width = sum(widths)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(total_width * 567)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    tbl_grid = table._tbl.find(qn("w:tblGrid"))
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        table._tbl.insert(1, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(int(width * 567)))
        tbl_grid.append(grid_col)

    for column, width in zip(table.columns, widths):
        column.width = Cm(width)


def apply_cell_basics(cell, width_cm: float) -> None:
    set_cell_width(cell, width_cm)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(4)


def set_document_defaults(document: Document) -> None:
    section = document.sections[0]
    section.orientation = WD_ORIENTATION.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    normal = document.styles["Normal"]
    normal.font.name = EAST_ASIA_FONT
    normal.font.size = Pt(10)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)


def add_title(document: Document, export_date: str = "") -> None:
    display_date = export_date.replace("-", ".")
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("交易文件主要条款摘要")
    title_run.bold = True
    title_run.font.name = TITLE_FONT
    title_run.font.size = Pt(12)
    title_run._element.rPr.rFonts.set(qn("w:eastAsia"), TITLE_FONT)

    date = document.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date.add_run(display_date)
    date_run.bold = True
    date_run.font.name = EAST_ASIA_FONT
    date_run.font.size = Pt(12)
    date_run.font.color.rgb = RGBColor(0, 0, 0)
    date_run._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)


def add_table_shell(document: Document, widths: list[float]):
    table = document.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    set_table_dimensions(table, widths)
    set_table_borders(table)

    header = table.rows[0]
    for cell, width, text in zip(header.cells, widths, ["#", "事项", "内容"]):
        apply_cell_basics(cell, width)
        set_cell_shading(cell, "A5C9EB")
        set_cell_text(cell, text, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    return table


def add_group_row(table, widths: list[float], group: str) -> None:
    group_row = table.add_row()
    merged = group_row.cells[0].merge(group_row.cells[-1])
    apply_cell_basics(merged, sum(widths))
    set_cell_shading(merged, "DAE9F7")
    set_cell_text(merged, group, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)


def add_item_rows(
    table,
    widths: list[float],
    label: str,
    content_lines: list[str],
    source_refs: list[str],
) -> None:
    lines = content_lines or [""]
    row = table.add_row()
    for cell, width in zip(row.cells, widths):
        apply_cell_basics(cell, width)
    set_cell_text(row.cells[0], "", align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(row.cells[1], label, align=WD_ALIGN_PARAGRAPH.LEFT)
    append_cell_lines(row.cells[2], [str(line) for line in lines])
    append_source_refs(row.cells[2], source_refs)


def add_kts_table(document: Document, items: list[dict[str, Any]]) -> None:
    widths = [0.74, 4.25, 19.61]
    table = add_table_shell(document, widths)
    current_group = ""
    for item in items:
        group = str(item["group"])
        if group != current_group:
            current_group = group
            add_group_row(table, widths, group)
        add_item_rows(
            table,
            widths,
            str(item["label"]),
            list(item["content_lines"]),
            list(item["source_refs"]),
        )


def build_kts_docx(record: dict[str, Any], export_date: str = "") -> bytes:
    items = export_items(record)
    document = Document()
    set_document_defaults(document)
    add_title(document, export_date)
    add_kts_table(document, items)

    output = BytesIO()
    document.save(output)
    return output.getvalue()
